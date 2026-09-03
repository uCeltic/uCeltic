"""One Search History entry, rendered as the Word document a reader keeps (#190).

The log rolls at 50 (ADR-0024), so this file is the only durable copy of a search — it
is laid out to be read, printed and cited, and it is built from the stored snapshot
alone, so it reproduces exactly what the entry holds and nothing that has happened to
the corpus since.
"""
import io
import math

from django.utils import timezone
from docx import Document
from docx.shared import Pt

# Kept here beside the builder rather than in the view: what this file is written as is
# the exporter's business, and the view only passes it on.
DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def match_percentage(score):
    """The stored dissimilarity as the similarity a reader understands (ADR-0024).

    Mirrors `client/src/history/matchPercentage.ts`, clamp included: the score has no
    upper bound in the store, and a matcher that one day returned more than 1 must not
    put a negative percentage in an exported document.

    Rounded half *up* rather than with `round()`, which rounds a half to even: the file
    has to agree with the entry as it was displayed, and JavaScript's `Math.round` — what
    the profile page rounds with — takes 88.5 to 89 where `round()` would take it to 88.
    """
    return f"{math.floor(max(0.0, 1 - score) * 100 + 0.5)}%"


def _parameters(entry):
    """The four knobs on one line each, in the order the popover stacks them.

    Under the names the user tuned them by on the Advanced Search popover, never the wire
    names the model stores them under: someone reading the file months later recognises
    "Match Length", not `window_size_ratio`. That one is also converted back to the
    percentage the slider showed — 1.3 was 130 on screen — and rounded, because the float
    that survives a round trip through JSON is 130.00000000000003.
    """
    return [
        f"Match Length: {round(entry.window_size_ratio * 100)}%",
        f"Precision: {entry.step_size}",
        f"Dissimilarity Score: {entry.dissimilarity_threshold:.2f}",
        f"Top K Results: {entry.top_k}",
    ]


def _when_searched(entry):
    """Written out in full rather than abbreviated: this is a citable date on a page that
    carries no other context about when it was made.

    The zone is named because it has to be. The profile page shows this instant in the
    reader's *browser* zone; a file is read anywhere, months later, so it states the zone
    it is written in rather than letting the two disagree silently.
    """
    return timezone.localtime(entry.created_at).strftime("%d %B %Y at %H:%M %Z")


def build_entry_document(entry):
    """The whole document for one entry: header, then each Version and its ranked hits."""
    document = Document()

    # The query is the title — it is what the reader is looking for when they open the
    # file, and quoted so a one-word search still reads as the thing that was searched.
    document.add_heading(f"Search: “{entry.query}”", level=0)
    document.add_paragraph(f"Searched {_when_searched(entry)}")
    for parameter in _parameters(entry):
        document.add_paragraph(parameter)

    covered = ", ".join(version["title"] for version in entry.versions)
    document.add_paragraph(f"Versions searched: {covered}")

    for version in entry.versions:
        document.add_heading(version["title"], level=1)
        hits = version["hits"]
        if not hits:
            # Kept, never dropped: a column that found nothing is part of what this
            # search was. A column that *errored* never reached the snapshot at all.
            document.add_paragraph("No matches.")
            continue
        for hit in hits:
            # No line number and no folio locator: a hit carries no usable line
            # reference for prose and no Manuscript Locator, so the passage and its
            # match percentage are the whole of what one hit can honestly say (#190).
            paragraph = document.add_paragraph(style="List Number")
            percentage = paragraph.add_run(f"{match_percentage(hit['score'])} match")
            percentage.bold = True
            paragraph.add_run(f" — {hit['snippet']}")

    for section in document.sections:
        section.left_margin = section.right_margin = Pt(72)

    return document


def entry_docx_stream(entry):
    """The document as the stream a response hands to the browser, rewound to the start."""
    buffer = io.BytesIO()
    build_entry_document(entry).save(buffer)
    buffer.seek(0)
    return buffer


def entry_filename(entry):
    """What the file lands in the reader's downloads as.

    Named for the moment searched rather than the query: a query can be any text at all
    — Old Irish, punctuation, an entire sentence — and the timestamp sorts a folder of
    exports into the order the searches were made. Seconds are included because two
    searches a minute apart is normal use.
    """
    stamped = timezone.localtime(entry.created_at).strftime("%Y-%m-%d-%H%M%S")
    return f"search-{stamped}.docx"
